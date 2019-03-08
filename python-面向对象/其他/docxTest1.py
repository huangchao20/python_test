#coding=utf-8

import docx
import os
from win32com import client as wc

def dealDocxFile(filename):
    if os.path.isfile(filename):
        print(filename)
        file = docx.Document(filename)
        print("段落数: ", str(len(file.paragraphs)))
        #输出段落编号及段落内容
        for i in range(len(file.paragraphs)):
            print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)

def DocxAndDocFile(filename):
    if os.path.isfile(filename):
        dfile = os.path.splitext(filename)
        if dfile[1] == ".docx" :         #处理docx文件
            dealDocxFile(filename)
        elif dfile[1] == ".doc":        #处理doc文件
            print("处理doc文件[%s]..." % filename)
            w = wc.Dispatch('Word.Application')
            doc = w.Documents.Open(filename)
            newfile = "".join((dfile[0], ".docx"))
            doc.SaveAs(newfile)
            dealDocxFile(newfile)

def printDocxInfo(dpath):
    Keywords = "安装操作"
    filetype = ['.doc', '.docx']
    print(dpath)
    if os.path.isdir(dpath):            #入参是文件夹
        for d in os.listdir(dpath):
            filename = os.path.join(dpath, d)
            m = os.path.splitext(d)
            if os.path.isfile(filename) and Keywords in m[0] and m[1] in filetype:
                DocxAndDocFile(filename)
    elif os.path.isfile(dpath):         #入参是文件
        m = os.path.splitext(dpath)
        if Keywords in m[0] and m[1] in filetype:
            DocxAndDocFile(dpath)
"""
from win32com import client as wc 
w = wc.Dispatch('Word.Application') 
# 或者使用下面的方法，使用启动独立的进程： 
# w = wc.DispatchEx('Word.Application') 
doc=w.Documents.Open("E:\\Jupyter\\s.doc") 
doc.SaveAs("E:\\Jupyter\\sa.docx",16)

#读取docx中的文本代码示例
import docx
#获取文档对象
file=docx.Document("D:\\temp\\word.docx")
print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段
 
#输出每一段的内容
for para in file.paragraphs:
 print(para.text)
 
#输出段落编号及段落内容
for i in range(len(file.paragraphs)):
 print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)
"""

if __name__ == '__main__':
    dpath = "E:\\SVN\\2019\\20190110\\特色业务平台\\fbap.20190110r.t10\\XQ-2018-796"
    printDocxInfo(dpath)